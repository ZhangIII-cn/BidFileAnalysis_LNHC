import docx
import pythoncom
import win32com
from win32com import client
import os
import zipfile
from oletools.olevba import VBA_Parser
import olefile
import openpyxl

Keyword_list=["无人机","巡检","联系人"]

def Doc_to_Docx(path):
    #利用win32com将doc转化为docx
    #print(path)
    #path='F:\Shell\File_Fliter\Test_doc\招标公告（2025年第4次服务）.doc'
    pythoncom.CoInitialize()
    word=win32com.client.Dispatch("Word.Application")
    #doc = word.Documents.Open(os.getcwd()+'/'+path)  #需要绝对路径
    #doc.SaveAs(os.getcwd()+'/'+path+"x",12)  #将doc后缀改成docx
    #-----------使用正确的绝对路径：
    # doc = word.Documents.Open(path)  # 需要绝对路径
    # doc.SaveAs(path+"x",12)  #将doc后缀改成docx
    # doc.Close()
    # word.Quit()
    try:
        doc = word.Documents.Open(path)  # 需要绝对路径
        doc.SaveAs(path+"x",12)  #将doc后缀改成docx
        doc.Close()
        word.Quit()
    except:
        print("Error：doc文件受损")
        return 3

def Detect_Package(path):
    with open(path, 'rb') as f:
        header = f.read(8)  # 读取前8字节
    # 常见文件头签名
    signatures = {
        b'%PDF-': "PDF",
        b'\x50\x4B\x03\x04': "ZIP (Office文件如 .docx/.xlsx)",
        b'\xD0\xCF\x11\xE0': "OLE (旧版Office如 .doc/.xls)",
        b'\x89PNG\r\n\x1A\n': "PNG图片",
        b'\xFF\xD8\xFF': "JPEG图片",
        b'\x47\x49\x46\x38': "GIF图片",
        b'\x25\x21\x50\x53': "PostScript (.ps)",
        b'\x7FELF': "ELF可执行文件",
        b'PK\x03\x04':"DOCX文件",
        b'\t\x08\x10':"XLS工作表",
        b'\xd0\xcf\x11\xe0':"DOC文件"
    }
    for sig, file_type in signatures.items():
        if header.startswith(sig):
            pass
            #print(file_type)



def Check_Inline_File(path):
    #个别文件中存在内嵌文件对象
    #基于docx本质是zip的原理，解析其embeddings目录下是否存在文件并打开解析
    #print(path)

    #----------------------------------利用oletools快速检查ole对象是否存在-----------------
    # 未解压情况下，oletools的方法无法探测到macros与olecontainer，原因待查
    # parser = VBA_Parser(path)
    # for filename, stream_data in parser.extract_macros():
    #     if filename.startswith('oleObject'):
    #         print(f"发现嵌入对象: {filename}")
    #         with open(filename, 'wb') as f:
    #             f.write(stream_data)


    #---------------------------------利用zip解压分析内嵌ole对象-----------------------------
    #在已知案例中，目标文件会以.bin形式保存于embadding目录
    #该文件前缀 D0 CF 11 E0 A1 B1 1A E1，为 application/vnd.visio(vsd) 格式
    #不排除后续可能存在直接嵌入docx和xlsx的情况
    try:
        with zipfile.ZipFile(path, 'r') as zip_ref:
            for file in zip_ref.namelist():
                if file.startswith("word/embeddings/") and file != "word/embeddings/":  #忽略掉文件夹目录
                    output_dir = os.path.abspath(os.path.dirname(path))  #解压文件输出目的目录的绝对路径
                    output_file = zip_ref.extract(file,output_dir)    #输出当前inline文件解压后保存的文件路径
                    #print(output_file)
                    File_strs = output_file.split('.')
                    File_EXT = File_strs[-1]
                    if File_EXT == 'bin' :
                        Figure_bin(output_file)
                    elif File_EXT == 'docx' or File_EXT == 'doc':
                        Figure_doc(output_file)
                    elif File_EXT == 'xlsx':
                        Figure_xls(output_file)
                    elif File_EXT == 'pdf':
                        Figure_pdf(output_file)
                    else :
                        print("存在未知格式文件："+output_file)
                        return 3

    except Exception as e:
        print(f"ZIP解压错误: {str(e)}")
        return 3

def Figure_bin(path):
    # 解析bin文件的原始类型 并尝试读取内容
    # 可能包含 doc xls docx xlsx pdf package等多种格式
    #print("This is bin!"+path)
    with olefile.OleFileIO(path) as ole:
        #print(ole.listdir())
        if ole.exists('Workbook'):    #存在工作表说明是xls文件,但是这是充分条件
            print("存在excel流")
            data = ole.openstream('Workbook').read()
            output_path=os.getcwd()+"/tmp_output/tmp.xls"
            with open(output_path, "wb") as f:
                f.write(data)
            Figure_xls(output_path)    #使用分析函数分析此xls

        elif ole.exists('package'):  #package存在多种可能
            print("存在package流")
            data=ole.openstream('package').read()
            #print(data)
            output_path = os.getcwd() + "/tmp_output/tmp2.doc"
            with open(output_path,'wb') as f:
                f.write(data)    #目前可以直接输出doc文件，但不确定是否存在其他状况
            Detect_Package(output_path)  #检查文件类型
        else:
            return 3


def Figure_doc(File_Path):
    print("-------------------------------------" + File_Path)

    EXT=File_Path.split('.')[-1]
    if EXT == 'doc':  #先将doc转化成docx
        if Doc_to_Docx(File_Path) == 3:
            return 3
        else :
            File_Path += "x"

    doc=docx.Document(File_Path)  #读取word文件
    paragraphs = doc.paragraphs   #处理文字内容
    for para in paragraphs:
        #print(para)
        if any(kw in para.text for kw in Keyword_list):
            return 1

    tables=doc.tables  #处理表格内容
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                if any(kw in cell.text for kw in Keyword_list):
                    return 1

    RE_Check_Inline_File = Check_Inline_File(File_Path)
    if RE_Check_Inline_File == 1 or RE_Check_Inline_File == 3:
        return RE_Check_Inline_File
    else:
        return 2



def Figure_xls(File_Path):
    #print(File_Path)
    pass



def Figure_pdf(File_Path):
    pass




# 返回值 1：符合筛选条件  2：不符合筛选条件  3：内部存在未知文件，需要人工核查
# if __name__ == "__main__":
#     Figure_doc('Test_doc/aaa.doc')



